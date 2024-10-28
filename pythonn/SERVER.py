import google.generativeai as generative_ai
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS  # Import CORS

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:5173"}})
# CORS(app, resources={r"/generate_report": {"origins": "http://192.168.0.162:3000/"}})

API_KEY = 'AIzaSyCCiKT6_gfxzXNEiYO3VeE_sY15h3pSf9Y'
generative_ai.configure(api_key=API_KEY)

# Function to call Google Generative AI
def ai(query):
    generative_ai.configure(api_key=API_KEY)
    model = generative_ai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(query)
    cleaned_response = response.text.replace('*', '').replace('#', '')
    return cleaned_response

# Fetch and paraphrase content from URL
def fetch_and_paraphrase(title, subtopic, sub_subtopic, url, length):
    try:
        # Get the content of the webpage
        response = requests.get(url)
        response.raise_for_status()  # Raise an error for bad responses
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract all paragraphs from the website
        paragraphs = soup.find_all('p')
        text_content = ' '.join([para.get_text() for para in paragraphs])

        # Filter content that mentions either the subtopic or sub-subtopic
        if sub_subtopic:
            related_paragraphs = [
                para.get_text() for para in paragraphs
                if subtopic.lower() in para.get_text().lower() or sub_subtopic.lower() in para.get_text().lower()
            ]
        else:
            related_paragraphs = [
                para.get_text() for para in paragraphs
                if subtopic.lower() in para.get_text().lower()
            ]

        if not related_paragraphs:
            return f"No relevant content found for '{subtopic}' and '{sub_subtopic}' on the website."

        # Combine filtered content into one text
        relevant_content = ' '.join(related_paragraphs)

        # Create paraphrasing prompt based on whether sub-subtopic exists
        if sub_subtopic:
            paraphrasing_prompt = (
                f"This is a section of a research paper or report. Please focus on the following details:\n\n"
                f"*Title*: {title}\n"
                f"*Subtopic*: {subtopic}\n"
                f"*Sub-subtopic*: {sub_subtopic}\n\n"
                f"Please summarize and paraphrase the key points from the following content, focusing on the "
                f"sub-subtopic '{sub_subtopic}' and the specific parent '{subtopic}' . "
                f"The final content should be approximately {length} words, maintaining clarity and relevance, while ignoring any irrelevant information:\n"
                f"{relevant_content}"
            )
        else:
            paraphrasing_prompt = (
                f"This is a section of a research paper or report. Please focus on the following details:\n\n"
                f"*Title*: {title}\n"
                f"*Subtopic*: {subtopic}\n\n"
                f"Please summarize and paraphrase the key points from the following content, focusing on the subtopic '{subtopic}'. "
                f"The final content should be approximately {length} words, maintaining clarity and relevance, while ignoring any irrelevant information:\n"
                f"{relevant_content}"
            )

        # Generate paraphrased content based on the prompt
        paraphrased_content = ai(paraphrasing_prompt)
        return paraphrased_content

    except Exception as e:
        return f"Error fetching data: {e}"



def create_report(user_data):
    doc = Document()

    # Add title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(user_data['Title'])
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "Prepared by"
    details = doc.add_paragraph()
    detailsrun1 = details.add_run("Prepared by: ")
    detailsrun1.font.bold = True
    detailsrun2 = details.add_run(f"{user_data['Name']}")

    toc_entries = []  # Store TOC entries
    references = []  # Store references

    # Placeholder for Table of Contents
    toc_placeholder = doc.add_paragraph("TOC_PLACEHOLDER")

    # Check if Subtopics is "AI generated" or manually input by the user
    if user_data['Subtopics']:
        # Non-AI generated: loop through subtopics and generate content
        for i, subtopic_data in enumerate(user_data['Subtopics'], 1):
            subtopic_title = subtopic_data['Title']
            toc_entries.append(f"{i}. {subtopic_title}")
            tit=doc.add_paragraph()
            tit.add_run(f"{i}. {subtopic_title}").bold=True
            


            # Generate or fetch content for the subtopic
            if subtopic_data['Link']:
                content = fetch_and_paraphrase(
                    title=user_data['Title'],
                    subtopic=subtopic_title,
                    sub_subtopic=None,
                    url=subtopic_data['Link'],
                    length=subtopic_data['Length']
                )
                references.append(subtopic_data['Link'])  # Add reference link
            else:
                content = ai(
                    f"This is a section of a research paper or report. Please focus on the following details:\n\n"
                    f"*Title*: {user_data['Title']}\n"
                    f"*Subtopic*: {subtopic_title}\n\n"
                    f"Generate approximately {subtopic_data['Length']} words of content focusing on the subtopic '{subtopic_title}'. "
                    f"The content should be clear, relevant, and informative, with attention to the main ideas related to the subtopic. "
                    f"Ensure that unnecessary or irrelevant information is excluded and the generated text fits within the scope of a report structure. give only content but DO NOT print the  Title again(i WANT CONTENT ONLY)"
                )

            doc.add_paragraph(content)

            # Handle sub-subtopics
            if 'SubSubtopics' in subtopic_data:
                for j, sub_subtopic_data in enumerate(subtopic_data['SubSubtopics'], 1):
                    sub_subtopic_title = sub_subtopic_data['Title']
                    toc_entries.append(f"   {i}.{j} {sub_subtopic_title}")
                    tit=doc.add_paragraph()
                    tit.add_run(f"{i}.{j} {sub_subtopic_title}").bold=True

                    # Generate or fetch content for the sub-subtopic
                    if sub_subtopic_data['Link']:
                        content = fetch_and_paraphrase(
                            title=user_data['Title'],
                            subtopic=subtopic_title,
                            sub_subtopic=sub_subtopic_title,
                            url=sub_subtopic_data['Link'],
                            length=sub_subtopic_data['Length']
                        )
                        references.append(sub_subtopic_data['Link'])  # Add reference link
                    else:
                        content = ai(
                            f"This is a section of a research paper or report. Please focus on the following details:\n\n"
                            f"*Title*: {user_data['Title']}\n"
                            f"*Subtopic*: {subtopic_title}\n"
                            f"*Sub-subtopic*: {sub_subtopic_title}\n\n"
                            f"Generate approximately {sub_subtopic_data['Length']} words of content focusing on the sub-subtopic '{sub_subtopic_title}', "
                            f"under the parent subtopic '{subtopic_title}'. The content should delve deeper into the sub-subtopic, providing specific and relevant details. "
                            f"Ensure that unnecessary or irrelevant information is excluded, while maintaining clarity and focus on the key points. give only content but DO NOT print the  Title again(i WANT CONTENT ONLY)"
                
                        )

                    doc.add_paragraph(content)
    else:
        # Entirely AI-generated content
        content = ai(
            f"This is a section of a research paper or report. Please focus on the following details:\n\n"
            f"*Title*: {user_data['Title']}\n\n"
            f"Generate approximately {user_data['Report Length']} words of content. "
            f"The content should be clear, relevant, and informative, fitting within the scope of the report."
        )
        doc.add_paragraph(content)

    # Add References at the end
    if references:
        doc.add_heading("References", level=1)
        for ref in references:
            doc.add_paragraph(ref)

    # Add TOC at the start, after title and "Prepared by"
    # Access and replace the TOC placeholder
    toc_paragraph = toc_placeholder  # Directly access the placeholder paragraph

    # Clear the placeholder by removing all runs
    for run in toc_paragraph.runs:
        run.clear()  # This will remove the placeholder text

    # Add the actual TOC heading and entries
    toc_paragraph.add_run("Table of Contents").bold = True

    # Add TOC entries
    for entry in toc_entries:
        toc_paragraph.add_run(f"\n{entry}").bold=True

    # Save document
    doc.save(f"{user_data['Title'].replace(' ', '_')}.docx")
    docuu=f"{user_data['Title'].replace(' ', '_')}.docx"
    return docuu




def extract_website_name(url):

    try:
        # Parse the URL and extract the netloc (network location)
        parsed_url = urlparse(url)
        # Split the netloc to get the website name
        website_name = parsed_url.netloc.split('.')[-2]  # Get the second-level domain (e.g., 'example' from 'example.com')
        return website_name.capitalize()  # Capitalize for better formatting
    except Exception as e:
        return f"Invalid URL: {url}"





# Function to format references based on the user's style
def format_reference(style, link, title):
    # Common citation style structures
    if style.lower() == "apa":
        # APA format: Title. Retrieved from URL
        formatted_reference = f"{title}. Retrieved from {link}"

    elif style.lower() == "mla":
        # MLA format: "Title." Website Name, Day Month Year, URL.
        formatted_reference = f"\"{title}.\" Retrieved from {link}."

    elif style.lower() == "chicago":
        # Chicago format: "Title of Webpage," Website Name, accessed Month Day, Year, URL.
        formatted_reference = f"\"{title},\" accessed from {link}."

    elif style.lower() == "harvard":
        # Harvard format: Title. (Year) Available at: URL (Accessed: Day Month Year).
        formatted_reference = f"{title}. Available at: {link}."

    elif style.lower() == "vancouver":
        # Vancouver format: Title [Internet]. Available from: URL.
        formatted_reference = f"{title} [Internet]. Available from: {link}."

    else:
        # For other or unknown styles, return a simple link and title format as fallback
        formatted_reference = f"{title}. Retrieved from {link}"

    return formatted_reference

# Example of formatting AI-suggested sources




def get_user_inputs():
    user_data = {}
    user_data['Name'] = input("Enter your name: ")
    user_data['Title'] = input("Enter the title of your report: ")
    ai_generated = input("Do you want a fully AI-generated report? (yes/no): ").lower()

    if ai_generated == 'no':
        subtopics = []
        subtopics_count = int(input("How many subtopics do you want in your outline?: "))

        for i in range(1, subtopics_count + 1):
            subtopic_title = input(f"Enter the title of subtopic {i}: ")
            subtopic_link = input(f"Enter the link for subtopic {subtopic_title} (if any, press enter to skip): ")
            subtopic_length = int(input(f"Enter the desired length (in words) for subtopic {subtopic_title}: "))

            sub_subtopics_count = int(input(f"How many sub-subtopics under {subtopic_title}?: "))
            sub_subtopics = []

            for j in range(1, sub_subtopics_count + 1):
                sub_subtopic_title = input(f"Enter the title of sub-subtopic {j}: ")
                sub_subtopic_link = input(f"Enter the link for sub-subtopic {sub_subtopic_title} (if any, press enter to skip): ")
                sub_subtopic_length = int(input(f"Enter the desired length (in words) for sub-subtopic {sub_subtopic_title}: "))

                sub_subtopics.append({
                    'Title': sub_subtopic_title,
                    'Link': sub_subtopic_link,
                    'Length': sub_subtopic_length
                })

            subtopics.append({
                'Title': subtopic_title,
                'Link': subtopic_link,
                'Length': subtopic_length,
                'SubSubtopics': sub_subtopics
            })
        user_data['Subtopics'] = subtopics
    else:
        user_data['Subtopics'] = []  # Empty list for AI-generated reports

    user_data['Report Length'] = int(input("Enter the length of the report (in words): "))

    tone_options = [
        "Descriptive",
        "Persuasive",
        "Professional",
        "Narrative",
        "Analytical",
        "Expository",
        "Research Paper",
        "Report Writing"
    ]

    print("Choose the tone of the report:")
    for i, option in enumerate(tone_options, 1):
        print(f"{i}. {option}")

    tone_choice = int(input("Select a tone (enter the number): ")) - 1
    user_data['Tone'] = tone_options[tone_choice]

    reference_styles = [
        "APA",
        "MLA",
        "Chicago",
        "Harvard",
        "IEEE",
        "Turabian",
        "Vancouver",
        "ACS"
    ]

    print("Choose the reference style:")
    for i, style in enumerate(reference_styles, 1):
        print(f"{i}. {style}")

    reference_choice = int(input("Select a reference style (enter the number): ")) - 1
    user_data['Reference Style'] = reference_styles[reference_choice]

    return user_data






# display_report_preferences(user_inputs)
@app.route('/generate_report', methods=['POST'])
def generate_report():
    user_data = request.get_json()  # Receiving the dictionary from React
    print("Recived")
    doc=create_report(user_data)
    print("Report Created")
    return send_file(doc,as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)


# {'Name': 'Jainil', 'Title': 'hell', 'Subtopics': [{'Title': 'Janil', 'Link': '', 'Length': 232, 'SubSubtopics': []}, {'Title': 'sahil', 'Link': '', 'Length': 1222, 'SubSubtopics': []}], 'Report Length': 111111, 'Tone': 'Descriptive', 'Reference Style': 'Harvard'}



# {'Name': 'Jainil', 'Title': 'Huffmann', 'Subtopics': [{'Title': 'History', 'Link': '', 'Length': 200, 'SubSubtopics': []}, {'Title': 'Code ', 'Link': '', 'Length': 300, 'SubSubtopics': [{'Title': 'Python', 'Link': '', 'Length': 200}, {'Title': 'Java', 'Link': '', 'Length': 200}]}, {'Title': 'Application', 'Link': '', 'Length': 300, 'SubSubtopics': [{'Title': 'Encryption', 'Link': '', 'Length': 200}, {'Title': 'Compression', 'Link': '', 'Length': 200}]}], 'Report Length': 300, 'Tone': 'Professional', 'Reference Style': 'IEEE'}